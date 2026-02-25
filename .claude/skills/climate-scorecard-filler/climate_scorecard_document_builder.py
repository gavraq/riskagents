"""
Climate Scorecard Document Builder

Generates professional Word documents for Climate & Environmental Risk Scorecards.
Includes branded cover page with www.risk-agents.com.

Aligned with PRA SS5/25 (December 2025) requirements.

Uses shared DocxBuilder from markdown-to-word skill for common document operations.
"""

from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import json
import sys

# Import will work when python-docx is installed: pip install python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word document generation will be limited.")
    print("Install with: pip install python-docx")

# Import shared DocxBuilder
try:
    sys.path.insert(0, str(Path(__file__).parent.parent / "markdown-to-word"))
    from docx_builder import DocxBuilder, BrandConfig
    SHARED_BUILDER_AVAILABLE = True
except ImportError:
    SHARED_BUILDER_AVAILABLE = False

# Import the scorecard classes
try:
    from climate_scorecard_helper import (
        EnhancedClimateScorecard,
        RiskAppetiteAlignment,
        ICaapConsiderations,
        LitigationRiskAssessment,
        ScenarioAnalysisQuality,
        DataQualityDeclaration
    )
except ImportError:
    pass  # Will work when called from the skill directory


class ClimateScorecardDocumentBuilder:
    """
    Builds professional Word documents for Climate & Environmental Risk Scorecards.

    Features:
    - Branded cover page with www.risk-agents.com
    - Professional formatting matching corporate standards
    - Full SS5/25 compliance sections
    - Tables for risk assessments
    - Color-coded risk indicators

    Uses shared DocxBuilder for common operations (margins, styles, tables,
    headings, footer) while retaining climate-specific assessment rendering.
    """

    # Brand colors (also available via shared BrandConfig)
    BRAND_PRIMARY = RGBColor(0, 51, 102) if DOCX_AVAILABLE else None
    BRAND_SECONDARY = RGBColor(0, 102, 153) if DOCX_AVAILABLE else None
    BRAND_ACCENT = RGBColor(0, 153, 204) if DOCX_AVAILABLE else None

    # Risk colors
    COLOR_LOW = RGBColor(0, 128, 0) if DOCX_AVAILABLE else None
    COLOR_MODERATE = RGBColor(255, 165, 0) if DOCX_AVAILABLE else None
    COLOR_HIGH = RGBColor(255, 0, 0) if DOCX_AVAILABLE else None
    COLOR_CRITICAL = RGBColor(139, 0, 0) if DOCX_AVAILABLE else None

    def __init__(self):
        self.document = None
        # Initialise shared builder with climate brand config
        if SHARED_BUILDER_AVAILABLE and DOCX_AVAILABLE:
            self._builder = DocxBuilder(brand=BrandConfig(
                primary=self.BRAND_PRIMARY,
                secondary=self.BRAND_SECONDARY,
                accent=self.BRAND_ACCENT,
                color_low=self.COLOR_LOW,
                color_moderate=self.COLOR_MODERATE,
                color_high=self.COLOR_HIGH,
                color_critical=self.COLOR_CRITICAL,
                org_name="Risk Agents",
                website="www.risk-agents.com",
            ))
        else:
            self._builder = None

    def create_scorecard_document(
        self,
        scorecard: 'EnhancedClimateScorecard',
        bank_name: str = "Example Bank",
        prepared_for: str = "Credit Committee"
    ) -> Document:
        """
        Create a professional Word document for the climate scorecard.

        Args:
            scorecard: EnhancedClimateScorecard object with all assessments
            bank_name: Name of the bank/institution
            prepared_for: Target audience (e.g., "Credit Committee", "Board Risk Committee")

        Returns:
            python-docx Document object
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

        self.document = Document()
        self._setup_document_styles()
        self._set_document_margins()

        # Add cover page
        self._add_cover_page(
            counterparty=scorecard.counterparty,
            bank_name=bank_name,
            prepared_for=prepared_for,
            assessment_date=scorecard.assessment_date,
            prepared_by=scorecard.prepared_by
        )

        # Add page break after cover
        self.document.add_page_break()

        # Executive Summary
        self._add_executive_summary(scorecard)

        # Counterparty Profile
        self._add_counterparty_profile(scorecard)

        # Transition Risk Assessment
        self._add_transition_risk_section(scorecard)

        # Physical Risk Assessment
        self._add_physical_risk_section(scorecard)

        # SS5/25 Sections
        if scorecard.scenario_analysis_quality:
            self._add_scenario_analysis_section(scorecard)

        if scorecard.litigation_risk and scorecard.litigation_risk.treatment == "distinct_channel":
            self._add_litigation_risk_section(scorecard)

        # Combined Assessment
        self._add_combined_assessment(scorecard)

        # SS5/25: Risk Appetite Alignment
        if scorecard.risk_appetite_alignment:
            self._add_risk_appetite_section(scorecard)

        # SS5/25: ICAAP Considerations
        if scorecard.icaap_considerations:
            self._add_icaap_section(scorecard)

        # SS5/25: Data Quality Declaration
        if scorecard.data_quality:
            self._add_data_quality_section(scorecard)

        # Conclusions and Recommendations
        self._add_conclusions_section(scorecard)

        # Monitoring Triggers
        self._add_monitoring_section(scorecard)

        # Regulatory References
        self._add_regulatory_references()

        # Footer with branding
        self._add_footer()

        return self.document

    def _setup_document_styles(self):
        """Setup custom document styles. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder._setup_styles()
        else:
            styles = self.document.styles
            try:
                heading_style = styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
                heading_style.font.size = Pt(14)
                heading_style.font.bold = True
                heading_style.font.color.rgb = self.BRAND_PRIMARY
            except:
                pass

    def _set_document_margins(self):
        """Set standard document margins. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder._set_margins()
        else:
            for section in self.document.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

    def _add_cover_page(
        self,
        counterparty: str,
        bank_name: str,
        prepared_for: str,
        assessment_date: str,
        prepared_by: str
    ):
        """Add branded cover page with www.risk-agents.com."""

        # Add spacing at top
        for _ in range(3):
            self.document.add_paragraph()

        # Website branding
        website = self.document.add_paragraph()
        website_run = website.add_run("www.risk-agents.com")
        website_run.font.size = Pt(14)
        website_run.font.color.rgb = self.BRAND_ACCENT
        website_run.font.italic = True
        website.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Main title
        title = self.document.add_paragraph()
        title_run = title.add_run("CLIMATE & ENVIRONMENTAL")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.BRAND_PRIMARY
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title2 = self.document.add_paragraph()
        title2_run = title2.add_run("RISK SCORECARD")
        title2_run.font.size = Pt(28)
        title2_run.font.bold = True
        title2_run.font.color.rgb = self.BRAND_PRIMARY
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        self.document.add_paragraph()

        # Regulatory alignment badge
        reg_badge = self.document.add_paragraph()
        reg_run = reg_badge.add_run("PRA SS5/25 Aligned")
        reg_run.font.size = Pt(12)
        reg_run.font.color.rgb = self.BRAND_SECONDARY
        reg_run.font.italic = True
        reg_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        for _ in range(3):
            self.document.add_paragraph()

        # Counterparty name (prominent)
        cp_para = self.document.add_paragraph()
        cp_run = cp_para.add_run(counterparty.upper())
        cp_run.font.size = Pt(24)
        cp_run.font.bold = True
        cp_run.font.color.rgb = self.BRAND_SECONDARY
        cp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        for _ in range(4):
            self.document.add_paragraph()

        # Document info table
        info_table = self.document.add_table(rows=5, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style the table
        for row in info_table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(3)

        # Populate table
        labels = ["Prepared For:", "Prepared By:", "Assessment Date:", "Bank/Institution:", "Document Version:"]
        values = [prepared_for, prepared_by or "Risk Management", assessment_date, bank_name, "3.0 (SS5/25)"]

        for i, (label, value) in enumerate(zip(labels, values)):
            cell0 = info_table.rows[i].cells[0]
            cell1 = info_table.rows[i].cells[1]

            run0 = cell0.paragraphs[0].add_run(label)
            run0.font.bold = True
            run0.font.size = Pt(11)
            run0.font.color.rgb = self.BRAND_PRIMARY

            run1 = cell1.paragraphs[0].add_run(value)
            run1.font.size = Pt(11)

        # Add spacing
        for _ in range(4):
            self.document.add_paragraph()

        # Confidentiality notice
        conf = self.document.add_paragraph()
        conf_run = conf.add_run("CONFIDENTIAL - For Internal Use Only")
        conf_run.font.size = Pt(10)
        conf_run.font.color.rgb = RGBColor(128, 128, 128)
        conf_run.font.italic = True
        conf.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generated by notice
        gen = self.document.add_paragraph()
        gen_run = gen.add_run("Generated by Risk Agents Platform")
        gen_run.font.size = Pt(9)
        gen_run.font.color.rgb = RGBColor(128, 128, 128)
        gen.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_executive_summary(self, scorecard: 'EnhancedClimateScorecard'):
        """Add executive summary section."""
        self._add_section_heading("Executive Summary")

        # Calculate scores
        overall_score = scorecard.calculate_overall_score()
        transition_score = scorecard.calculate_transition_score()
        physical_score = scorecard.calculate_physical_score()
        override_req, override_dir, override_notches = scorecard.requires_rating_override()

        # Risk level determination
        risk_level, risk_color = self._get_risk_level_and_color(overall_score)

        # Summary paragraph
        summary_para = self.document.add_paragraph()
        summary_para.add_run(f"{scorecard.counterparty} ").bold = True
        summary_para.add_run(f"demonstrates ")

        risk_run = summary_para.add_run(f"{risk_level} ")
        risk_run.bold = True
        risk_run.font.color.rgb = risk_color

        summary_para.add_run(f"climate and environmental risk (Overall Score: {overall_score:.1f}/5).")

        # Score summary table
        self.document.add_paragraph()
        score_table = self.document.add_table(rows=5, cols=3)
        score_table.style = 'Light Grid Accent 1'

        # Headers
        headers = ["Risk Category", "Score", "Rating"]
        for i, header in enumerate(headers):
            score_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

        # Data rows
        data = [
            ("Transition Risk", f"{transition_score:.1f}/5", self._score_to_rating(transition_score)),
            ("Physical Risk", f"{physical_score:.1f}/5", self._score_to_rating(physical_score)),
            ("Overall Climate Risk", f"{overall_score:.1f}/5", self._score_to_rating(overall_score)),
            ("Rating Override", "Yes" if override_req else "No", f"{override_notches} notch(es) {override_dir}" if override_req else "N/A")
        ]

        for i, (category, score, rating) in enumerate(data, 1):
            score_table.rows[i].cells[0].text = category
            score_table.rows[i].cells[1].text = score
            score_table.rows[i].cells[2].text = rating

        # Risk appetite alignment (if available)
        if scorecard.risk_appetite_alignment:
            self.document.add_paragraph()
            ra_para = self.document.add_paragraph()
            ra_para.add_run("Risk Appetite Alignment: ").bold = True
            ra_para.add_run(f"{scorecard.risk_appetite_alignment.category}")

        self.document.add_paragraph()

    def _add_counterparty_profile(self, scorecard: 'EnhancedClimateScorecard'):
        """Add counterparty profile section."""
        self._add_section_heading("Counterparty Profile")

        profile_table = self.document.add_table(rows=4, cols=2)
        profile_table.style = 'Light Grid Accent 1'

        data = [
            ("Counterparty", scorecard.counterparty),
            ("Country", scorecard.country),
            ("Sector", scorecard.sector),
            ("Sector Classification", scorecard.sector_classification or "Not Specified")
        ]

        for i, (label, value) in enumerate(data):
            profile_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
            profile_table.rows[i].cells[1].text = value

        self.document.add_paragraph()

    def _add_transition_risk_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add transition risk assessment section."""
        self._add_section_heading("Transition Risk Assessment")

        # Preparedness
        if scorecard.transition_preparedness:
            self._add_subsection_heading("1. Transition Preparedness (Intent & Capability)")

            prep = scorecard.transition_preparedness
            prep_table = self.document.add_table(rows=6, cols=3)
            prep_table.style = 'Light Grid Accent 1'

            # Headers
            for i, header in enumerate(["Assessment Factor", "Score", "Status"]):
                prep_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

            # Data
            factors = [
                ("Net-Zero Target & Credibility", prep.net_zero_target),
                ("Climate Disclosure Quality (TCFD/CDP)", prep.tcfd_disclosure),
                ("Governance Structure", prep.governance_structure),
                ("Transition Plan Credibility", prep.transition_plan),
                ("Capex Alignment with Paris Goals", prep.capex_alignment)
            ]

            for i, (factor, score) in enumerate(factors, 1):
                prep_table.rows[i].cells[0].text = factor
                prep_table.rows[i].cells[1].text = f"{score}/5" if score else "N/A"
                prep_table.rows[i].cells[2].text = self._score_to_rating(score) if score else "Not Assessed"

            # Rationale
            if prep.rationale:
                self.document.add_paragraph()
                rat_para = self.document.add_paragraph()
                rat_para.add_run("Rationale: ").bold = True
                rat_para.add_run(prep.rationale)

            self.document.add_paragraph()

        # Vulnerability
        if scorecard.transition_vulnerability:
            self._add_subsection_heading("2. Transition Vulnerability (Exposure & Risk)")

            vuln = scorecard.transition_vulnerability
            vuln_table = self.document.add_table(rows=8, cols=3)
            vuln_table.style = 'Light Grid Accent 1'

            # Headers
            for i, header in enumerate(["Risk Factor", "Score", "Level"]):
                vuln_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

            # Data
            factors = [
                ("Sector Carbon Intensity", vuln.sector_carbon_intensity),
                ("Stranded Asset Risk", vuln.stranded_asset_risk),
                ("Policy & Regulatory Pressure", vuln.policy_regulatory_risk),
                ("Technology Disruption Vulnerability", vuln.technology_disruption),
                ("Market & Investor Sentiment Risk", vuln.market_sentiment_risk),
                ("Legal & Litigation Risk", vuln.legal_litigation_risk),
                ("Country Transition Dependency", vuln.country_transition_dependency)
            ]

            for i, (factor, score) in enumerate(factors, 1):
                vuln_table.rows[i].cells[0].text = factor
                vuln_table.rows[i].cells[1].text = f"{score}/5" if score else "N/A"
                vuln_table.rows[i].cells[2].text = self._score_to_rating(score) if score else "Not Assessed"

            if vuln.rationale:
                self.document.add_paragraph()
                rat_para = self.document.add_paragraph()
                rat_para.add_run("Rationale: ").bold = True
                rat_para.add_run(vuln.rationale)

            self.document.add_paragraph()

        # Opportunity
        if scorecard.transition_opportunity:
            self._add_subsection_heading("3. Transition Opportunity Assessment")

            opp = scorecard.transition_opportunity
            opp_table = self.document.add_table(rows=4, cols=3)
            opp_table.style = 'Light Grid Accent 1'

            # Headers
            for i, header in enumerate(["Opportunity Factor", "Score", "Positioning"]):
                opp_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

            factors = [
                ("Low-Carbon Market Growth Potential", opp.market_growth_potential),
                ("Green Revenue Share", opp.green_revenue_share),
                ("Competitive Advantage in Transition", opp.competitive_advantage)
            ]

            for i, (factor, score) in enumerate(factors, 1):
                opp_table.rows[i].cells[0].text = factor
                opp_table.rows[i].cells[1].text = f"{score}/5" if score else "N/A"
                opp_table.rows[i].cells[2].text = self._opportunity_label(score) if score else "Not Assessed"

            self.document.add_paragraph()

    def _add_physical_risk_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add physical risk assessment section."""
        self._add_section_heading("Physical Risk Assessment")

        if scorecard.physical_risk:
            phys = scorecard.physical_risk
            phys_table = self.document.add_table(rows=6, cols=3)
            phys_table.style = 'Light Grid Accent 1'

            # Headers
            for i, header in enumerate(["Risk Factor", "Score", "Level"]):
                phys_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

            factors = [
                ("Acute Climate Hazard Exposure", phys.acute_hazard_exposure),
                ("Chronic Climate Change Exposure", phys.chronic_climate_exposure),
                ("Ecosystem Dependency & Vulnerability", phys.ecosystem_dependency),
                ("Adaptation Capability", phys.adaptation_capability),
                ("Physical Risk Scenario Analysis", phys.scenario_analysis_done)
            ]

            for i, (factor, score) in enumerate(factors, 1):
                phys_table.rows[i].cells[0].text = factor
                phys_table.rows[i].cells[1].text = f"{score}/5" if score else "N/A"
                phys_table.rows[i].cells[2].text = self._score_to_rating(score) if score else "Not Assessed"

            if phys.rationale:
                self.document.add_paragraph()
                rat_para = self.document.add_paragraph()
                rat_para.add_run("Rationale: ").bold = True
                rat_para.add_run(phys.rationale)

        self.document.add_paragraph()

    def _add_scenario_analysis_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add SS5/25 scenario analysis quality section."""
        self._add_section_heading("Scenario Analysis Quality (SS5/25)")

        note = self.document.add_paragraph()
        note.add_run("SS5/25 requires assessment of how CSA results inform actual decision-making").italic = True

        self.document.add_paragraph()

        saq = scorecard.scenario_analysis_quality
        saq_table = self.document.add_table(rows=5, cols=3)
        saq_table.style = 'Light Grid Accent 1'

        for i, header in enumerate(["Factor", "Score", "Assessment"]):
            saq_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

        factors = [
            ("Analysis Conducted", saq.analysis_conducted),
            ("Results Integrated into Strategy", saq.results_integrated),
            ("Time Horizons Appropriate", saq.horizons_appropriate),
            ("Documentation Quality", saq.documentation_quality)
        ]

        for i, (factor, score) in enumerate(factors, 1):
            saq_table.rows[i].cells[0].text = factor
            saq_table.rows[i].cells[1].text = f"{score}/5" if score else "N/A"
            saq_table.rows[i].cells[2].text = self._score_to_rating(score) if score else "Not Assessed"

        self.document.add_paragraph()

    def _add_litigation_risk_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add SS5/25 litigation risk section (distinct channel)."""
        self._add_section_heading("Litigation Risk Assessment (SS5/25 Distinct Channel)")

        note = self.document.add_paragraph()
        note.add_run("Per SS5/25, litigation risk assessed as distinct transmission channel").italic = True

        self.document.add_paragraph()

        lit = scorecard.litigation_risk
        lit_table = self.document.add_table(rows=4, cols=2)
        lit_table.style = 'Light Grid Accent 1'

        data = [
            ("Treatment", "Distinct Channel"),
            ("Litigation Type", lit.litigation_type or "Not specified"),
            ("Score", f"{lit.score}/5" if lit.score else "N/A"),
            ("Rationale", lit.rationale or "Not provided")
        ]

        for i, (label, value) in enumerate(data):
            lit_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
            lit_table.rows[i].cells[1].text = value

        self.document.add_paragraph()

    def _add_combined_assessment(self, scorecard: 'EnhancedClimateScorecard'):
        """Add combined assessment section."""
        self._add_section_heading("Combined Assessment")

        transition_score = scorecard.calculate_transition_score()
        physical_score = scorecard.calculate_physical_score()
        overall_score = scorecard.calculate_overall_score()
        override_req, override_dir, override_notches = scorecard.requires_rating_override()

        combined_table = self.document.add_table(rows=7, cols=2)
        combined_table.style = 'Light Grid Accent 1'

        data = [
            ("Transition Risk Score", f"{transition_score:.1f}/5"),
            ("Physical Risk Score", f"{physical_score:.1f}/5"),
            ("Overall Climate Risk Score", f"{overall_score:.1f}/5"),
            ("Sector Classification", scorecard.sector_classification or "None"),
            ("Rating Override Required", "Yes" if override_req else "No"),
            ("Override Direction", override_dir),
            ("Override Notches", str(override_notches) if override_req else "N/A")
        ]

        for i, (label, value) in enumerate(data):
            combined_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
            combined_table.rows[i].cells[1].text = value

        self.document.add_paragraph()

    def _add_risk_appetite_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add SS5/25 risk appetite alignment section."""
        self._add_section_heading("Risk Appetite Alignment (SS5/25)")

        note = self.document.add_paragraph()
        note.add_run("SS5/25 requires explicit alignment with firm's climate risk appetite framework").italic = True

        self.document.add_paragraph()

        ra = scorecard.risk_appetite_alignment
        ra_table = self.document.add_table(rows=4, cols=2)
        ra_table.style = 'Light Grid Accent 1'

        data = [
            ("Appetite Category", ra.category),
            ("Justification", ra.justification),
            ("Portfolio Limit Impact", ra.portfolio_limit_impact or "N/A"),
            ("Escalation Required", f"Yes - {ra.escalation_reason}" if ra.escalation_required else "No")
        ]

        for i, (label, value) in enumerate(data):
            ra_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
            ra_table.rows[i].cells[1].text = value

        self.document.add_paragraph()

    def _add_icaap_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add SS5/25 ICAAP considerations section."""
        self._add_section_heading("Capital & ICAAP Considerations (SS5/25 - Banks)")

        note = self.document.add_paragraph()
        note.add_run("SS5/25 requires banks to evidence climate risk materiality in ICAAPs").italic = True

        self.document.add_paragraph()

        ic = scorecard.icaap_considerations
        ic_table = self.document.add_table(rows=3, cols=2)
        ic_table.style = 'Light Grid Accent 1'

        data = [
            ("Capital Relevance", ic.capital_relevance),
            ("ICAAP Treatment", ic.icaap_treatment),
            ("Materiality Justification", ic.materiality_justification)
        ]

        for i, (label, value) in enumerate(data):
            ic_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
            ic_table.rows[i].cells[1].text = value

        self.document.add_paragraph()

    def _add_data_quality_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add SS5/25 data quality declaration section."""
        self._add_section_heading("Data Quality Declaration (SS5/25)")

        note = self.document.add_paragraph()
        note.add_run("SS5/25 requires documentation of data sources, proxies, and limitations").italic = True

        self.document.add_paragraph()

        dq = scorecard.data_quality

        # Primary sources
        self._add_subsection_heading("Primary Data Sources")
        for source in dq.primary_sources:
            self.document.add_paragraph(f"• {source}", style='List Bullet')

        # Proxies used
        if dq.proxies_used:
            self._add_subsection_heading("Proxies Used")
            proxy_table = self.document.add_table(rows=len(dq.proxies_used) + 1, cols=3)
            proxy_table.style = 'Light Grid Accent 1'

            for i, header in enumerate(["Data Point", "Proxy Applied", "Limitation"]):
                proxy_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

            for i, proxy in enumerate(dq.proxies_used, 1):
                proxy_table.rows[i].cells[0].text = proxy.get('data_point', 'N/A')
                proxy_table.rows[i].cells[1].text = proxy.get('proxy', 'N/A')
                proxy_table.rows[i].cells[2].text = proxy.get('limitation', 'N/A')

        # Data gaps
        if dq.data_gaps:
            self._add_subsection_heading("Identified Data Gaps")
            for gap in dq.data_gaps:
                self.document.add_paragraph(f"• {gap}", style='List Bullet')

        # Uncertainty acknowledgment
        if dq.uncertainty_acknowledgment:
            self._add_subsection_heading("Uncertainty Acknowledgment")
            self.document.add_paragraph(dq.uncertainty_acknowledgment)

        self.document.add_paragraph()

    def _add_conclusions_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add conclusions and recommendations section."""
        self._add_section_heading("Conclusions & Recommendations")

        if scorecard.summary:
            self.document.add_paragraph(scorecard.summary)

        # Key risk drivers
        if scorecard.key_risk_drivers:
            self._add_subsection_heading("Key Risk Drivers")
            for driver in scorecard.key_risk_drivers:
                self.document.add_paragraph(f"• {driver}", style='List Bullet')

        # Key opportunities
        if scorecard.key_opportunities:
            self._add_subsection_heading("Key Opportunities")
            for opp in scorecard.key_opportunities:
                self.document.add_paragraph(f"• {opp}", style='List Bullet')

        # Mitigation measures
        if scorecard.mitigation_measures:
            self._add_subsection_heading("Mitigation Measures in Place")
            for measure in scorecard.mitigation_measures:
                self.document.add_paragraph(f"• {measure}", style='List Bullet')

        self.document.add_paragraph()

    def _add_monitoring_section(self, scorecard: 'EnhancedClimateScorecard'):
        """Add monitoring triggers section."""
        self._add_section_heading("Monitoring & Review Triggers")

        if scorecard.monitoring_triggers:
            for trigger in scorecard.monitoring_triggers:
                self.document.add_paragraph(f"• {trigger}", style='List Bullet')

        self.document.add_paragraph()

    def _add_regulatory_references(self):
        """Add regulatory references section."""
        self._add_section_heading("Regulatory References")

        refs = [
            "PRA SS5/25 - Enhancing banks' and insurers' approaches to managing climate-related risks (December 2025)",
            "BCBS Principles - Principles for effective management and supervision of climate-related financial risks",
            "ISSB IFRS S2 - Climate-related Disclosures"
        ]

        for ref in refs:
            self.document.add_paragraph(f"• {ref}", style='List Bullet')

        self.document.add_paragraph()

    def _add_footer(self):
        """Add footer with branding. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder.add_footer()
        else:
            footer_para = self.document.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = footer_para.add_run("─" * 50)
            run.font.color.rgb = RGBColor(200, 200, 200)

            self.document.add_paragraph()

            gen_para = self.document.add_paragraph()
            gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            gen_run = gen_para.add_run(f"Generated by Risk Agents Platform | www.risk-agents.com | {datetime.now().strftime('%d %B %Y')}")
            gen_run.font.size = Pt(9)
            gen_run.font.color.rgb = RGBColor(128, 128, 128)
            gen_run.font.italic = True

    def _add_section_heading(self, title: str):
        """Add a section heading. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder.add_section_heading(title)
        else:
            heading = self.document.add_heading(title, level=1)
            for run in heading.runs:
                run.font.color.rgb = self.BRAND_PRIMARY

    def _add_subsection_heading(self, title: str):
        """Add a subsection heading. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder.add_subsection_heading(title)
        else:
            heading = self.document.add_heading(title, level=2)
            for run in heading.runs:
                run.font.color.rgb = self.BRAND_SECONDARY

    def _get_risk_level_and_color(self, score: float) -> tuple:
        """Get risk level text and color based on score. Uses shared builder if available."""
        if self._builder:
            return self._builder.get_risk_level_text(score), self._builder.get_risk_color(score)
        if score <= 1.5:
            return "NEGLIGIBLE", self.COLOR_LOW
        elif score <= 2.5:
            return "LOW", self.COLOR_LOW
        elif score <= 3.5:
            return "MODERATE", self.COLOR_MODERATE
        elif score <= 4.5:
            return "HIGH", self.COLOR_HIGH
        else:
            return "CRITICAL", self.COLOR_CRITICAL

    def _score_to_rating(self, score: Optional[int]) -> str:
        """Convert score to rating label. Uses shared builder if available."""
        if self._builder:
            return self._builder.score_to_rating(score)
        if score is None:
            return "Not Assessed"
        elif score <= 1.5:
            return "Excellent"
        elif score <= 2.5:
            return "Good"
        elif score <= 3.5:
            return "Adequate"
        elif score <= 4.5:
            return "Weak"
        else:
            return "Critical"

    def _opportunity_label(self, score: Optional[int]) -> str:
        """Convert opportunity score to label."""
        if score is None:
            return "Not Assessed"
        elif score <= 1.5:
            return "Strong Beneficiary"
        elif score <= 2.5:
            return "Some Benefit"
        elif score <= 3.5:
            return "Neutral"
        elif score <= 4.5:
            return "Limited Benefit"
        else:
            return "No Benefit"

    def save_document(self, output_path: Path) -> Path:
        """Save document to file."""
        if self.document:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(str(output_path))
            return output_path
        else:
            raise RuntimeError("No document to save")


# Convenience function for direct use
def generate_climate_scorecard_docx(
    scorecard: 'EnhancedClimateScorecard',
    output_path: str,
    bank_name: str = "Example Bank",
    prepared_for: str = "Credit Committee"
) -> Path:
    """
    Generate a Word document from a climate scorecard.

    Args:
        scorecard: EnhancedClimateScorecard object
        output_path: Path for output .docx file
        bank_name: Name of the bank/institution
        prepared_for: Target audience

    Returns:
        Path to saved document
    """
    builder = ClimateScorecardDocumentBuilder()
    builder.create_scorecard_document(scorecard, bank_name, prepared_for)
    return builder.save_document(Path(output_path))


# Example usage
if __name__ == "__main__":
    from climate_scorecard_helper import EnhancedClimateScorecard

    # Create example scorecard
    scorecard = EnhancedClimateScorecard(
        counterparty="Example Energy Corp",
        country="United Kingdom",
        sector="Energy",
        prepared_by="Risk Management Team"
    )

    # Add assessments
    scorecard.assess_transition_preparedness(
        net_zero_target=3,
        tcfd_disclosure=2,
        governance_structure=3,
        transition_plan=3,
        capex_alignment=4,
        rationale="Moderate transition preparedness with room for improvement"
    )

    scorecard.assess_transition_vulnerability(
        sector_carbon_intensity=4,
        stranded_asset_risk=4,
        policy_regulatory_risk=3,
        technology_disruption=3,
        market_sentiment_risk=3,
        legal_litigation_risk=2,
        country_transition_dependency=2,
        rationale="High vulnerability due to sector exposure"
    )

    scorecard.assess_physical_risk(
        acute_hazard_exposure=2,
        chronic_climate_exposure=3,
        ecosystem_dependency=2,
        adaptation_capability=3,
        scenario_analysis_done=4,
        rationale="Moderate physical risk exposure"
    )

    # SS5/25 sections
    scorecard.set_risk_appetite_alignment(
        category="Manage",
        justification="Within monitored energy sector appetite",
        portfolio_limit_impact="Counts toward 15% energy sector limit"
    )

    scorecard.set_icaap_considerations(
        capital_relevance="Medium",
        icaap_treatment="Stress Testing",
        materiality_justification="Material energy sector exposure"
    )

    scorecard.set_data_quality_declaration(
        primary_sources=["Annual Report 2024", "CDP Submission", "TCFD Report"],
        proxies_used=[{"data_point": "Scope 3", "proxy": "Sector average", "limitation": "May underestimate"}],
        data_gaps=["No independent verification of transition plan"],
        uncertainty_acknowledgment="Transition timeline uncertain due to policy changes"
    )

    scorecard.summary = "Example Energy Corp demonstrates MODERATE climate risk requiring active monitoring."
    scorecard.key_risk_drivers = ["High carbon intensity", "Stranded asset exposure"]
    scorecard.mitigation_measures = ["TCFD disclosure", "Net-zero commitment"]
    scorecard.monitoring_triggers = ["Annual review", "Policy changes"]

    # Generate document
    builder = ClimateScorecardDocumentBuilder()
    doc = builder.create_scorecard_document(
        scorecard,
        bank_name="Example Bank",
        prepared_for="Credit Committee"
    )
    builder.save_document(Path("output/example_climate_scorecard.docx"))
    print("Document generated successfully!")
