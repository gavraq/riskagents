"""
DocxBuilder - Shared Word Document Generation Utilities

Provides reusable document building operations used across Risk Agents skills
that generate Word (.docx) output. Extracts common patterns from
MLRCDocumentBuilder and ClimateScorecardDocumentBuilder.

Usage:
    from docx_builder import DocxBuilder, BrandConfig

    builder = DocxBuilder()
    builder.create_document()
    builder.add_cover_page(title="Report Title")
    builder.add_section_heading("Section 1")
    builder.add_paragraph("Content here")
    builder.save("output/report.docx")
"""

from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─── Brand Configuration ──────────────────────────────────────────────────────

@dataclass
class BrandConfig:
    """Configurable brand colours and metadata for document generation."""

    # Colours
    primary: Any = None       # Main heading colour
    secondary: Any = None     # Sub-heading colour
    accent: Any = None        # Accent/highlight colour

    # Risk level colours
    color_low: Any = None
    color_moderate: Any = None
    color_high: Any = None
    color_critical: Any = None

    # Organisation details
    org_name: str = "Risk Agents"
    website: str = "www.risk-agents.com"
    confidentiality: str = "CONFIDENTIAL - For Internal Use Only"

    def __post_init__(self):
        """Set defaults using RGBColor if available."""
        if DOCX_AVAILABLE:
            if self.primary is None:
                self.primary = RGBColor(0, 51, 102)
            if self.secondary is None:
                self.secondary = RGBColor(0, 102, 153)
            if self.accent is None:
                self.accent = RGBColor(0, 153, 204)
            if self.color_low is None:
                self.color_low = RGBColor(0, 128, 0)
            if self.color_moderate is None:
                self.color_moderate = RGBColor(255, 165, 0)
            if self.color_high is None:
                self.color_high = RGBColor(255, 0, 0)
            if self.color_critical is None:
                self.color_critical = RGBColor(139, 0, 0)


# ─── Default brand config ─────────────────────────────────────────────────────

DEFAULT_BRAND = BrandConfig()


# ─── DocxBuilder ───────────────────────────────────────────────────────────────

class DocxBuilder:
    """
    Shared Word document builder with common utilities.

    Provides standard operations for creating professional Word documents
    with consistent formatting, branded cover pages, and structured tables.
    """

    def __init__(self, brand: Optional[BrandConfig] = None):
        """
        Initialise the builder.

        Args:
            brand: Optional brand configuration. Uses Risk Agents defaults if not provided.
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        self.brand = brand or DEFAULT_BRAND
        self.document: Optional[Document] = None

    # ─── Document Lifecycle ────────────────────────────────────────────────

    def create_document(self) -> Document:
        """Create a new document with standard setup."""
        self.document = Document()
        self._set_margins()
        self._setup_styles()
        return self.document

    def save(self, output_path: str | Path) -> Path:
        """Save document to file. Creates parent directories if needed."""
        if not self.document:
            raise RuntimeError("No document to save. Call create_document() first.")
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(path))
        return path

    # ─── Document Setup ────────────────────────────────────────────────────

    def _set_margins(
        self,
        top: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        right: float = 1.0,
    ):
        """Set document margins in inches."""
        for section in self.document.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)

    def _setup_styles(self):
        """Setup custom document styles."""
        styles = self.document.styles
        try:
            heading_style = styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.color.rgb = self.brand.primary
        except ValueError:
            pass  # Style already exists

    # ─── Cover Page ────────────────────────────────────────────────────────

    def add_cover_page(
        self,
        title: str,
        subtitle: str = "",
        prepared_for: str = "",
        prepared_by: str = "",
        date: Optional[str] = None,
        regulatory_badge: str = "",
        extra_fields: Optional[List[Tuple[str, str]]] = None,
    ):
        """
        Add a branded cover page.

        Args:
            title: Main document title
            subtitle: Subtitle or document type
            prepared_for: Target audience
            prepared_by: Author/team
            date: Document date (defaults to today)
            regulatory_badge: Optional regulatory alignment text (e.g., "PRA SS5/25 Aligned")
            extra_fields: Additional label/value pairs for the info table
        """
        date = date or datetime.now().strftime("%d %B %Y")

        # Top spacing
        for _ in range(3):
            self.document.add_paragraph()

        # Website branding
        website_para = self.document.add_paragraph()
        website_run = website_para.add_run(self.brand.website)
        website_run.font.size = Pt(14)
        website_run.font.color.rgb = self.brand.accent
        website_run.font.italic = True
        website_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Main title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.brand.primary
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            sub_para = self.document.add_paragraph()
            sub_run = sub_para.add_run(subtitle.upper())
            sub_run.font.size = Pt(20)
            sub_run.font.bold = True
            sub_run.font.color.rgb = self.brand.primary
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Regulatory badge
        if regulatory_badge:
            self.document.add_paragraph()
            badge_para = self.document.add_paragraph()
            badge_run = badge_para.add_run(regulatory_badge)
            badge_run.font.size = Pt(12)
            badge_run.font.color.rgb = self.brand.secondary
            badge_run.font.italic = True
            badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing before info table
        for _ in range(4):
            self.document.add_paragraph()

        # Info table
        info_fields = []
        if prepared_for:
            info_fields.append(("Prepared For:", prepared_for))
        if prepared_by:
            info_fields.append(("Prepared By:", prepared_by))
        info_fields.append(("Date:", date))
        info_fields.append(("Organisation:", self.brand.org_name))
        if extra_fields:
            info_fields.extend(extra_fields)

        if info_fields:
            table = self.document.add_table(rows=len(info_fields), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row in table.rows:
                row.cells[0].width = Inches(2)
                row.cells[1].width = Inches(3)

            for i, (label, value) in enumerate(info_fields):
                run0 = table.rows[i].cells[0].paragraphs[0].add_run(label)
                run0.font.bold = True
                run0.font.size = Pt(11)
                run0.font.color.rgb = self.brand.primary

                run1 = table.rows[i].cells[1].paragraphs[0].add_run(value)
                run1.font.size = Pt(11)

        # Bottom spacing
        for _ in range(4):
            self.document.add_paragraph()

        # Confidentiality notice
        conf_para = self.document.add_paragraph()
        conf_run = conf_para.add_run(self.brand.confidentiality)
        conf_run.font.size = Pt(10)
        conf_run.font.color.rgb = RGBColor(128, 128, 128)
        conf_run.font.italic = True
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generated by
        gen_para = self.document.add_paragraph()
        gen_run = gen_para.add_run(f"Generated by {self.brand.org_name} Platform")
        gen_run.font.size = Pt(9)
        gen_run.font.color.rgb = RGBColor(128, 128, 128)
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break after cover
        self.document.add_page_break()

    # ─── Headings ──────────────────────────────────────────────────────────

    def add_section_heading(self, title: str, level: int = 1):
        """Add a section heading with brand colour."""
        heading = self.document.add_heading(title, level=level)
        for run in heading.runs:
            run.font.color.rgb = self.brand.primary

    def add_subsection_heading(self, title: str, level: int = 2):
        """Add a subsection heading with secondary brand colour."""
        heading = self.document.add_heading(title, level=level)
        for run in heading.runs:
            run.font.color.rgb = self.brand.secondary

    def add_title_heading(self, title: str):
        """Add a centred title heading (level 0)."""
        heading = self.document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── Paragraphs ────────────────────────────────────────────────────────

    def add_paragraph(self, text: str = "", bold: bool = False, italic: bool = False):
        """Add a paragraph with optional formatting."""
        para = self.document.add_paragraph()
        if text:
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
        return para

    def add_bold_text(self, label: str, value: str = ""):
        """Add a paragraph with bold label followed by normal text."""
        para = self.document.add_paragraph()
        para.add_run(label).bold = True
        if value:
            para.add_run(f" {value}")
        return para

    def add_bullet_list(self, items: List[str]):
        """Add a bulleted list."""
        for item in items:
            self.document.add_paragraph(f"• {item}", style='List Bullet')

    def add_centred_text(
        self, text: str, size: int = 11, color: Optional[Any] = None, italic: bool = False
    ):
        """Add centred text with optional formatting."""
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.font.italic = italic
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para

    def add_date_line(self, date: Optional[str] = None):
        """Add a centred date line."""
        date_text = date or datetime.now().strftime("%d %B %Y")
        self.add_centred_text(date_text)

    def add_spacing(self, lines: int = 1):
        """Add empty paragraphs for spacing."""
        for _ in range(lines):
            self.document.add_paragraph()

    def add_page_break(self):
        """Add a page break."""
        self.document.add_page_break()

    # ─── Tables ────────────────────────────────────────────────────────────

    def add_key_value_table(
        self,
        data: List[Tuple[str, str]],
        style: str = 'Light Grid Accent 1',
    ):
        """
        Add a two-column key-value table.

        Args:
            data: List of (label, value) tuples
            style: Word table style name
        """
        table = self.document.add_table(rows=len(data), cols=2)
        table.style = style

        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
            table.rows[i].cells[1].text = str(value)

        self.add_spacing()
        return table

    def add_data_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        style: str = 'Light Grid Accent 1',
        bold_headers: bool = True,
    ):
        """
        Add a multi-column data table with headers.

        Args:
            headers: Column header texts
            rows: List of row data (each row is a list of cell values)
            style: Word table style name
            bold_headers: Whether to bold the header row
        """
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = style

        # Headers
        for i, header in enumerate(headers):
            run = table.rows[0].cells[i].paragraphs[0].add_run(header)
            run.bold = bold_headers

        # Data rows
        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = str(value)

        self.add_spacing()
        return table

    def add_cover_sheet_table(
        self,
        data: List[Tuple[str, str]],
        style: str = 'Light Grid Accent 1',
    ):
        """Add a cover sheet / board report table (used in MLRC documents)."""
        return self.add_key_value_table(data, style)

    # ─── Risk Level Utilities ──────────────────────────────────────────────

    def get_risk_color(self, score: float) -> Any:
        """Get colour for a risk score (1-5 scale)."""
        if score <= 1.5:
            return self.brand.color_low
        elif score <= 2.5:
            return self.brand.color_low
        elif score <= 3.5:
            return self.brand.color_moderate
        elif score <= 4.5:
            return self.brand.color_high
        else:
            return self.brand.color_critical

    def get_risk_level_text(self, score: float) -> str:
        """Get risk level text for a score (1-5 scale)."""
        if score <= 1.5:
            return "NEGLIGIBLE"
        elif score <= 2.5:
            return "LOW"
        elif score <= 3.5:
            return "MODERATE"
        elif score <= 4.5:
            return "HIGH"
        else:
            return "CRITICAL"

    def score_to_rating(self, score: Optional[int]) -> str:
        """Convert score to rating label (Excellent/Good/Adequate/Weak/Critical)."""
        if score is None:
            return "Not Assessed"
        elif score <= 1.5:
            return "Excellent"
        elif score <= 2.5:
            return "Good"
        elif score <= 3.5:
            return "Adequate"
        elif score <= 4.5:
            return "Weak"
        else:
            return "Critical"

    def add_colored_risk_text(self, label: str, score: float, para=None):
        """Add text with risk-appropriate colour."""
        if para is None:
            para = self.document.add_paragraph()
        run = para.add_run(label)
        run.bold = True
        run.font.color.rgb = self.get_risk_color(score)
        return para

    # ─── Footer ────────────────────────────────────────────────────────────

    def add_footer(self):
        """Add a standard footer with branding."""
        # Divider
        div_para = self.document.add_paragraph()
        div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        div_run = div_para.add_run("─" * 50)
        div_run.font.color.rgb = RGBColor(200, 200, 200)

        self.document.add_paragraph()

        # Generated by
        gen_para = self.document.add_paragraph()
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_run = gen_para.add_run(
            f"Generated by {self.brand.org_name} Platform | "
            f"{self.brand.website} | "
            f"{datetime.now().strftime('%d %B %Y')}"
        )
        gen_run.font.size = Pt(9)
        gen_run.font.color.rgb = RGBColor(128, 128, 128)
        gen_run.font.italic = True

    # ─── Document History / Governance ─────────────────────────────────────

    def add_document_history(
        self,
        prepared_by: str,
        date: str,
        reviewed_by: str = "N/A",
    ):
        """Add a standard document history section."""
        self.add_section_heading("Document History")
        self.add_key_value_table([
            ("Prepared by", prepared_by),
            ("Date", date),
            ("Reviewed by 2nd line of defence", reviewed_by),
        ])

    def add_formal_governance(
        self,
        committee: str = "MLRC",
        outcome: str = "*(To be completed post-meeting)*",
        matters: str = "*(To be completed post-meeting)*",
    ):
        """Add a formal governance section."""
        self.add_section_heading("Formal Document Governance")
        self.add_key_value_table([
            ("Reviewing committee and meeting date", committee),
            ("Outcome and key rationale for decision", outcome),
            ("Significant matters raised and associated actions", matters),
        ])

    def add_references(self, refs: Optional[List[str]] = None):
        """Add a references section."""
        self.add_section_heading("References")
        if refs:
            self.add_bullet_list(refs)
        else:
            self.add_paragraph("[1] See attached reference documents")
        self.add_spacing()

    # ─── Markdown Conversion ───────────────────────────────────────────────

    def add_markdown_content(self, markdown_text: str):
        """
        Convert markdown to Word paragraphs with proper formatting.

        Supports:
        - Headings (# through ####)
        - **Bold** and *italic* inline formatting (including nested)
        - Bullet lists (- or *)
        - Numbered lists (1. through 999.)
        - Markdown pipe tables (| col1 | col2 |)
        - Horizontal rules (--- or ***)
        - Page breaks (---pagebreak--- or \\pagebreak)
        """
        import re

        lines = markdown_text.strip().split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Headings (#### through #)
            if stripped.startswith('#'):
                level = 0
                for ch in stripped:
                    if ch == '#':
                        level += 1
                    else:
                        break
                heading_text = stripped[level:].strip()
                if level == 1:
                    self.add_section_heading(heading_text, level=1)
                elif level == 2:
                    self.add_subsection_heading(heading_text, level=2)
                elif level >= 3:
                    heading = self.document.add_heading(heading_text, level=min(level, 4))
                    for run in heading.runs:
                        run.font.color.rgb = self.brand.secondary
                i += 1
                continue

            # Markdown table detection (line starts with | and contains |)
            if stripped.startswith('|') and '|' in stripped[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_markdown_table(table_lines)
                continue

            # Bullet lists (- or * at start, but not ---)
            if (stripped.startswith('- ') or stripped.startswith('* ')) and not stripped.startswith('---'):
                text = stripped[2:]
                para = self.document.add_paragraph(style='List Bullet')
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Nested bullet lists (indented with spaces/tabs)
            if re.match(r'^(\s{2,}|\t+)[-*] ', line):
                text = re.sub(r'^[\s\t]+[-*] ', '', line)
                para = self.document.add_paragraph(style='List Bullet 2')
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Numbered lists (1. through 999.)
            num_match = re.match(r'^(\d{1,3})\.\s+(.*)', stripped)
            if num_match:
                text = num_match.group(2)
                para = self.document.add_paragraph(style='List Number')
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Horizontal rule / page break
            if stripped in ('---', '***', '___', '---pagebreak---', '\\pagebreak'):
                if 'pagebreak' in stripped:
                    self.document.add_page_break()
                else:
                    self.document.add_paragraph()
                i += 1
                continue

            # Regular paragraph with inline formatting
            para = self.document.add_paragraph()
            self._add_formatted_runs(para, stripped)
            i += 1

    def _add_formatted_runs(self, paragraph, text: str):
        """
        Parse inline markdown formatting and add as Word runs with proper styles.

        Handles: **bold**, *italic*, ***bold italic***, `code`,
        and combinations like **bold *and italic* text**.
        """
        import re

        if not text:
            return

        # Pattern matches: ***bold italic***, **bold**, *italic*, `code`, or plain text
        # Order matters: match *** before ** before *
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'   # ***bold italic***
            r'|(\*\*(.+?)\*\*)'       # **bold**
            r'|(\*(.+?)\*)'           # *italic*
            r'|(`(.+?)`)'             # `code`
            r'|([^*`]+)'              # plain text
        )

        for match in pattern.finditer(text):
            if match.group(2):  # ***bold italic***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(4):  # **bold**
                run = paragraph.add_run(match.group(4))
                run.bold = True
            elif match.group(6):  # *italic*
                run = paragraph.add_run(match.group(6))
                run.italic = True
            elif match.group(8):  # `code`
                run = paragraph.add_run(match.group(8))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 0, 0)
            elif match.group(9):  # plain text
                paragraph.add_run(match.group(9))

    def _add_markdown_table(self, table_lines: list[str]):
        """
        Convert markdown pipe table lines to a Word table.

        Handles:
        - Header row (first line)
        - Separator row (|---|---|, skipped)
        - Data rows
        """
        if not table_lines:
            return

        def parse_row(line: str) -> list[str]:
            """Parse a markdown table row into cell values."""
            # Strip leading/trailing pipes and split
            line = line.strip()
            if line.startswith('|'):
                line = line[1:]
            if line.endswith('|'):
                line = line[:-1]
            return [cell.strip() for cell in line.split('|')]

        def is_separator(line: str) -> bool:
            """Check if line is a table separator (|---|---|)."""
            import re
            cleaned = line.replace('|', '').replace('-', '').replace(':', '').strip()
            return len(cleaned) == 0

        # Parse all rows, skipping separator lines
        parsed_rows = []
        for line in table_lines:
            if is_separator(line):
                continue
            parsed_rows.append(parse_row(line))

        if not parsed_rows:
            return

        # First row is headers
        headers = parsed_rows[0]
        data_rows = parsed_rows[1:]

        # Ensure all rows have the same number of columns
        num_cols = len(headers)
        for row in data_rows:
            while len(row) < num_cols:
                row.append('')

        # Create Word table
        table = self.document.add_table(rows=len(data_rows) + 1, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Header row
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(9)

        # Data rows
        for row_idx, row_data in enumerate(data_rows):
            for col_idx, value in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.paragraphs[0].clear()
                    # Apply inline formatting within table cells
                    self._add_formatted_runs(cell.paragraphs[0], value)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)

        self.add_spacing()

    # ─── Cell Shading Utility ──────────────────────────────────────────────

    @staticmethod
    def set_cell_shading(cell, color_hex: str):
        """
        Set background shading on a table cell.

        Args:
            cell: python-docx table cell
            color_hex: Hex colour string (e.g., "FF0000" for red)
        """
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
