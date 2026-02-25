"""
MLRC Document Builder

Generates Word documents in MLRC governance format for stress scenario documentation.
Supports both new scenario specifications and annual review memos.

Uses shared DocxBuilder from markdown-to-word skill for common document operations.
"""

from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import json
import sys

# Import will work when python-docx is installed: pip install python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word document generation will be limited.")
    print("Install with: pip install python-docx")

# Import shared DocxBuilder
try:
    sys.path.insert(0, str(Path(__file__).parent.parent / "markdown-to-word"))
    from docx_builder import DocxBuilder, BrandConfig
    SHARED_BUILDER_AVAILABLE = True
except ImportError:
    SHARED_BUILDER_AVAILABLE = False

from scenario_designer import StressScenario, AssetClassShocks
from scenario_reviewer import ScenarioReviewResult, AssetClassReview

class MLRCDocumentBuilder:
    """
    Builds MLRC-formatted Word documents for stress scenarios.

    Uses shared DocxBuilder for common operations (margins, headings, tables,
    document history, governance sections) while retaining MLRC-specific
    logic for scenario content, validation, and consultation sections.
    """

    def __init__(self):
        self.document = None
        self._builder = DocxBuilder() if SHARED_BUILDER_AVAILABLE and DOCX_AVAILABLE else None

    def create_new_scenario_document(
        self,
        scenario: StressScenario,
        meeting_date: str,
        presenter: str = "Market Risk"
    ) -> Document:
        """
        Create MLRC document for a new stress scenario.

        Args:
            scenario: StressScenario object
            meeting_date: MLRC meeting date
            presenter: Name of presenter

        Returns:
            python-docx Document object
        """
        if not DOCX_AVAILABLE:
            return self._create_markdown_fallback_new(scenario, meeting_date, presenter)

        self.document = Document()

        # Set document properties
        self._set_document_margins()

        # Add header
        self._add_header(scenario.metadata.scenario_name)

        # Cover Sheet
        self._add_cover_sheet_new(
            scenario_name=scenario.metadata.scenario_name,
            meeting_date=meeting_date,
            presenter=presenter
        )

        # Background / Economic Narrative
        self._add_background_section_new(scenario.narrative)

        # Asset Class Sections
        for asset_class_shock in scenario.asset_class_shocks:
            self._add_asset_class_section(asset_class_shock)

        # Validation Results (if any warnings/errors)
        if scenario.validation_results:
            self._add_validation_section(scenario.validation_results)

        # Front Office Consultation
        if scenario.consultation_prompts:
            self._add_consultation_section(scenario.consultation_prompts)

        # Conclusions
        self._add_conclusions_new(scenario)

        # References
        self._add_references()

        # Document History
        self._add_document_history(
            prepared_by=scenario.metadata.created_by,
            date=scenario.metadata.created_date
        )

        # Formal Governance
        self._add_formal_governance()

        return self.document

    def create_annual_review_document(
        self,
        review_result: ScenarioReviewResult,
        meeting_date: str
    ) -> Document:
        """
        Create MLRC document for annual scenario review.

        Args:
            review_result: ScenarioReviewResult object
            meeting_date: MLRC meeting date

        Returns:
            python-docx Document object
        """
        if not DOCX_AVAILABLE:
            return self._create_markdown_fallback_review(review_result, meeting_date)

        self.document = Document()
        self._set_document_margins()

        # Add header
        self._add_header(f"{review_result.scenario_name} - Annual Review {datetime.now().year}")

        # Cover Sheet
        self._add_cover_sheet_review(
            scenario_name=review_result.scenario_name,
            meeting_date=meeting_date,
            presenter=review_result.reviewer
        )

        # Background
        self._add_background_section_review(review_result)

        # Asset Class Reviews
        for asset_class_review in review_result.asset_class_reviews:
            self._add_asset_class_review_section(asset_class_review)

        # Conclusions
        self._add_conclusions_review(review_result)

        # References
        self._add_references()

        # Document History
        self._add_document_history(
            prepared_by=review_result.reviewer,
            date=review_result.review_date
        )

        # Formal Governance
        self._add_formal_governance()

        return self.document

    def _set_document_margins(self):
        """Set standard document margins. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder._set_margins()
        else:
            sections = self.document.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

    def _add_header(self, title: str):
        """Add document header with logo placeholder."""
        header = self.document.add_heading(title, level=0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_para = self.document.add_paragraph(datetime.now().strftime("%d %B %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()  # Spacing

    def _add_cover_sheet_new(self, scenario_name: str, meeting_date: str, presenter: str):
        """Add cover sheet for new scenario."""
        self.document.add_heading("Board Report Cover Sheet", level=1)

        # Create table for cover sheet
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Document Name
        table.rows[0].cells[0].text = "Document Name"
        table.rows[0].cells[1].text = f"{scenario_name} - Scenario Parameters {datetime.now().year}"

        # Meeting Date
        table.rows[1].cells[0].text = "Meeting Date"
        table.rows[1].cells[1].text = f"MLRC - {meeting_date}"

        # Presenter
        table.rows[2].cells[0].text = "Presenter"
        table.rows[2].cells[1].text = presenter

        # Context
        table.rows[3].cells[0].text = "Context"
        table.rows[3].cells[1].text = f"New pillar stress scenario for approval - {scenario_name}"

        # Purpose
        table.rows[4].cells[0].text = "Purpose"
        table.rows[4].cells[1].text = f"Approval of new stress scenario parameterization for {scenario_name}"

        self.document.add_paragraph()  # Spacing

    def _add_cover_sheet_review(self, scenario_name: str, meeting_date: str, presenter: str):
        """Add cover sheet for annual review."""
        self.document.add_heading("Board Report Cover Sheet", level=1)

        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = "Document Name"
        table.rows[0].cells[1].text = f"{scenario_name} - Annual Review {datetime.now().year}"

        table.rows[1].cells[0].text = "Meeting Date"
        table.rows[1].cells[1].text = f"MLRC - {meeting_date}"

        table.rows[2].cells[0].text = "Presenter"
        table.rows[2].cells[1].text = presenter

        table.rows[3].cells[0].text = "Context"
        table.rows[3].cells[1].text = f"Memo for noting - Annual Review for the \"{scenario_name}\" Scenario"

        table.rows[4].cells[0].text = "Purpose"
        table.rows[4].cells[1].text = f"Annual review of \"{scenario_name}\" scenario shocks applied to various asset classes"

        self.document.add_paragraph()

    def _add_background_section_new(self, narrative):
        """Add background section for new scenario."""
        self.document.add_heading("Background", level=1)

        # Trigger Event
        self.document.add_paragraph().add_run("Trigger Event: ").bold = True
        self.document.add_paragraph(narrative.trigger_event)

        # Narrative Text
        self.document.add_paragraph(narrative.narrative_text)

        # Key Assumptions
        if narrative.key_assumptions:
            self.document.add_paragraph().add_run("Key Assumptions:").bold = True
            for assumption in narrative.key_assumptions:
                self.document.add_paragraph(f"• {assumption}", style='List Bullet')

        self.document.add_paragraph()

    def _add_background_section_review(self, review_result: ScenarioReviewResult):
        """Add background section for annual review."""
        self.document.add_heading("Background", level=1)

        para = self.document.add_paragraph()
        para.add_run(f"The below are the shocks approved for the \"{review_result.scenario_name}\" scenario as per the Example Bank Stress Testing Parameterisation document.")

        if review_result.market_changes_since_last_review:
            self.document.add_paragraph().add_run("Market Developments Since Last Review:").bold = True
            for change in review_result.market_changes_since_last_review:
                self.document.add_paragraph(f"• {change}", style='List Bullet')

        if review_result.system_changes:
            self.document.add_paragraph().add_run("System and Platform Changes:").bold = True
            for change in review_result.system_changes:
                self.document.add_paragraph(f"• {change}", style='List Bullet')

        self.document.add_paragraph()

    def _add_asset_class_section(self, asset_class_shock: AssetClassShocks):
        """Add asset class section with shock specifications."""
        self.document.add_heading(asset_class_shock.asset_class, level=1)

        # Description
        self.document.add_paragraph(asset_class_shock.description)

        # Rationale
        self.document.add_paragraph().add_run("Rationale: ").bold = True
        self.document.add_paragraph(asset_class_shock.rationale)

        # Historical Analogue
        if asset_class_shock.historical_analogue:
            para = self.document.add_paragraph()
            para.add_run("Historical Analogue: ").bold = True
            para.add_run(asset_class_shock.historical_analogue).italic = True

        # Shocks (simplified representation - full implementation would create detailed tables)
        self.document.add_paragraph().add_run("Proposed Shocks:").bold = True
        self._add_shock_table(asset_class_shock.shocks, asset_class_shock.asset_class)

        self.document.add_paragraph()

    def _add_shock_table(self, shocks: Dict, asset_class: str):
        """Add table with shock specifications (simplified)."""
        # This is a simplified version - full implementation would create
        # detailed tables matching the exact format from the Word documents

        para = self.document.add_paragraph(f"*(Detailed {asset_class} shock specifications - see attached Excel parameterization file)*")
        para.italic = True

        # For demonstration, show key shocks
        for key, value in list(shocks.items())[:5]:  # Show first 5
            self.document.add_paragraph(f"• {key}: {value}", style='List Bullet')

    def _add_asset_class_review_section(self, review: AssetClassReview):
        """Add asset class review section."""
        self.document.add_heading(review.asset_class, level=1)

        if review.change_type.value == "no_change":
            para = self.document.add_paragraph()
            para.add_run("No proposed changes").bold = True
            self.document.add_paragraph()
            self.document.add_paragraph("Current Applied Shocks")
            self.document.add_paragraph("*(See attached parameterization document)*").italic = True
        else:
            para = self.document.add_paragraph()
            para.add_run(review.change_type.value.replace('_', ' ').title()).bold = True

            self.document.add_paragraph().add_run("Rationale: ").bold = True
            self.document.add_paragraph(review.rationale)

            if review.proposed_parameters:
                self.document.add_paragraph().add_run("Proposed Changes:").bold = True
                for param, value in review.proposed_parameters.items():
                    self.document.add_paragraph(f"• {param}: {value}", style='List Bullet')

            if review.affected_products:
                para = self.document.add_paragraph()
                para.add_run("Affected Products: ").bold = True
                para.add_run(", ".join(review.affected_products))

        self.document.add_paragraph()

    def _add_validation_section(self, validation_results):
        """Add validation results section."""
        from validators import ValidationLevel

        errors = [v for v in validation_results if v.level == ValidationLevel.ERROR]
        warnings = [v for v in validation_results if v.level == ValidationLevel.WARNING]

        if errors or warnings:
            self.document.add_heading("Validation Results", level=1)

            if errors:
                self.document.add_paragraph().add_run("ERRORS (Must Address):").bold = True
                for error in errors:
                    para = self.document.add_paragraph(f"• {error.message}", style='List Bullet')
                    para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            if warnings:
                self.document.add_paragraph().add_run("WARNINGS (Review Recommended):").bold = True
                for warning in warnings:
                    para = self.document.add_paragraph(f"• {warning.message}", style='List Bullet')
                    para.runs[0].font.color.rgb = RGBColor(255, 165, 0)

            self.document.add_paragraph()

    def _add_consultation_section(self, consultation_prompts: Dict):
        """Add Front Office consultation section."""
        self.document.add_heading("Front Office Consultation", level=1)

        for desk, prompts in consultation_prompts.items():
            self.document.add_paragraph().add_run(f"{desk}:").bold = True
            for prompt in prompts:
                self.document.add_paragraph(f"• {prompt}", style='List Bullet')

        self.document.add_paragraph()

    def _add_conclusions_new(self, scenario: StressScenario):
        """Add conclusions for new scenario."""
        self.document.add_heading("Conclusions and Recommended Actions", level=1)

        para = self.document.add_paragraph()
        para.add_run(f"Scenario submitted for MLRC approval. ").bold = True
        para.add_run(f"Confidence Score: {scenario.confidence_score:.0f}%")

        if scenario.next_steps:
            self.document.add_paragraph().add_run("Next Steps:").bold = True
            for step in scenario.next_steps:
                self.document.add_paragraph(f"• {step}", style='List Bullet')

        self.document.add_paragraph()

    def _add_conclusions_review(self, review_result: ScenarioReviewResult):
        """Add conclusions for annual review."""
        self.document.add_heading("Conclusions and Recommended Actions", level=1)

        action_text = {
            "approve_as_is": "Memo submitted for approval. No changes to scenario parameters proposed.",
            "approve_with_changes": "Memo submitted for approval with proposed parameter updates as detailed above.",
            "major_revision_needed": "Significant revision required. Recommend scheduling detailed scenario redesign discussion."
        }

        self.document.add_paragraph(action_text[review_result.recommended_action])
        self.document.add_paragraph()

    def _add_references(self):
        """Add references section. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder.add_references(["Example Bank Stress Testing Parameterisation"])
        else:
            self.document.add_heading("References", level=1)
            self.document.add_paragraph("[1] Example Bank Stress Testing Parameterisation")
            self.document.add_paragraph()

    def _add_document_history(self, prepared_by: str, date: str):
        """Add document history section. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder.add_document_history(prepared_by, date)
        else:
            self.document.add_heading("Document History", level=1)

            table = self.document.add_table(rows=3, cols=2)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = "Prepared by"
            table.rows[0].cells[1].text = prepared_by

            table.rows[1].cells[0].text = "Date"
            table.rows[1].cells[1].text = date

            table.rows[2].cells[0].text = "Reviewed by 2nd line of defence"
            table.rows[2].cells[1].text = "N/A"

            self.document.add_paragraph()

    def _add_formal_governance(self):
        """Add formal governance section. Delegates to shared builder if available."""
        if self._builder:
            self._builder.document = self.document
            self._builder.add_formal_governance(
                committee="MLRC",
                outcome="*(To be completed post-MLRC)*",
                matters="*(To be completed post-MLRC)*"
            )
        else:
            self.document.add_heading("Formal Document Governance", level=1)

            table = self.document.add_table(rows=3, cols=2)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = "Reviewing committee and meeting date"
            table.rows[0].cells[1].text = "MLRC"

            table.rows[1].cells[0].text = "Outcome and key rationale for decision"
            table.rows[1].cells[1].text = "*(To be completed post-MLRC)*"

            table.rows[2].cells[0].text = "Significant matters raised and associated actions"
            table.rows[2].cells[1].text = "*(To be completed post-MLRC)*"

    def save_document(self, output_path: Path) -> Path:
        """Save document to file."""
        if self.document and DOCX_AVAILABLE:
            self.document.save(str(output_path))
            return output_path
        else:
            raise RuntimeError("No document to save or python-docx not available")

    # Fallback methods for when python-docx is not installed

    def _create_markdown_fallback_new(self, scenario, meeting_date, presenter) -> str:
        """Create markdown version of new scenario document (fallback)."""
        md_content = scenario_designer.ScenarioDesigner._generate_markdown_output(scenario)
        return md_content

    def _create_markdown_fallback_review(self, review_result, meeting_date) -> str:
        """Create markdown version of review document (fallback)."""
        from scenario_reviewer import ScenarioReviewer
        reviewer = ScenarioReviewer(Path("data"))
        return reviewer.generate_review_memo(review_result)
